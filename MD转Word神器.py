#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MD转Word神器
功能：将Markdown文件转换为格式专业的Word文档
"""

import os
import sys
import re
from datetime import datetime

# 检查并导入必要的库
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
except ImportError:
    print("=" * 60)
    print("❌ 缺少必要的库：python-docx")
    print("=" * 60)
    print("请运行以下命令安装：")
    print("pip install python-docx")
    print("=" * 60)
    sys.exit(1)


class MarkdownToWordConverter:
    """Markdown转Word转换器"""
    
    def __init__(self):
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档基础样式"""
        # 设置页面边距
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
        
        # 设置默认样式
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.5
    
    def _is_work_log_or_audit_log(self, line):
        """检查是否是Work Log或Audit Log"""
        return line.strip().startswith('[Work Log]') or line.strip().startswith('[Audit Log]')
    
    def _add_paragraph_with_format(self, text, is_heading=False, heading_level=1, is_special=False):
        """添加段落并设置格式"""
        if is_heading:
            # 标题
            heading = self.doc.add_heading(text, level=heading_level)
            # 设置标题字体
            for run in heading.runs:
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
        else:
            # 正文段落
            para = self.doc.add_paragraph()
            
            # 处理特殊标记（Work Log / Audit Log）
            if is_special:
                run = para.add_run(text)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            else:
                # 处理中英文混排
                parts = self._split_chinese_english(text)
                for part in parts:
                    run = para.add_run(part)
                    if self._is_chinese(part):
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    else:
                        run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            
            # 设置段落格式
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    
    def _split_chinese_english(self, text):
        """分离中英文"""
        parts = []
        current = ""
        is_chinese = None
        
        for char in text:
            char_is_chinese = self._is_chinese_char(char)
            
            if is_chinese is None:
                is_chinese = char_is_chinese
                current = char
            elif is_chinese == char_is_chinese:
                current += char
            else:
                parts.append(current)
                current = char
                is_chinese = char_is_chinese
        
        if current:
            parts.append(current)
        
        return parts if parts else [text]
    
    def _is_chinese_char(self, char):
        """判断是否是中文字符"""
        return '\u4e00' <= char <= '\u9fff'
    
    def _is_chinese(self, text):
        """判断文本是否主要是中文"""
        if not text.strip():
            return True
        chinese_count = sum(1 for c in text if self._is_chinese_char(c))
        return chinese_count > len(text) * 0.3
    
    def _process_markdown_line(self, line):
        """处理Markdown行"""
        original_line = line
        line = line.rstrip()
        
        # 空行
        if not line.strip():
            return None
        
        # 标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            # 移除标题中的格式标记
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            return {'type': 'heading', 'level': level, 'text': text}
        
        # 特殊标记（Work Log / Audit Log）
        if self._is_work_log_or_audit_log(line):
            return {'type': 'special', 'text': line}
        
        # 加粗文本（**文本**）- 如果整行都是加粗或者是明显的加粗格式
        if line.strip().startswith('**') and line.strip().endswith('**'):
            return {'type': 'bold', 'text': line}
        
        # 普通段落
        return {'type': 'paragraph', 'text': original_line}
    
    def _process_bold_text(self, text):
        """处理加粗文本"""
        para = self.doc.add_paragraph()
        # 使用更精确的正则表达式匹配加粗文本
        parts = re.split(r'(\*\*[^*]+\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                # 加粗文本
                bold_text = part[2:-2]
                parts_bold = self._split_chinese_english(bold_text)
                for p in parts_bold:
                    run = para.add_run(p)
                    if self._is_chinese(p):
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    else:
                        run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = True
            else:
                # 普通文本
                if part.strip():
                    parts_normal = self._split_chinese_english(part)
                    for p in parts_normal:
                        run = para.add_run(p)
                        if self._is_chinese(p):
                            run.font.name = '宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        else:
                            run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
        
        # 设置段落格式
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.first_line_indent = Cm(0.74)
    
    def convert(self, markdown_file):
        """转换Markdown文件为Word文档"""
        if not os.path.exists(markdown_file):
            print(f"❌ 文件不存在：{markdown_file}")
            return None
        
        print(f"📖 正在读取文件：{markdown_file}")
        
        with open(markdown_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        print(f"✅ 文件读取成功，共 {len(lines)} 行")
        print("🔄 正在转换...")
        
        for line in lines:
            result = self._process_markdown_line(line)
            
            if result is None:
                # 空行，跳过
                continue
            
            if result['type'] == 'heading':
                self._add_paragraph_with_format(
                    result['text'],
                    is_heading=True,
                    heading_level=result['level']
                )
            elif result['type'] == 'special':
                self._add_paragraph_with_format(
                    result['text'],
                    is_special=True
                )
            elif result['type'] == 'bold':
                self._process_bold_text(result['text'])
            elif result['type'] == 'paragraph':
                # 检查是否包含引号内的逐字稿内容
                stripped = line.strip()
                if stripped.startswith('"') and stripped.endswith('"') and len(stripped) > 2:
                    # 逐字稿内容，特殊处理（去掉引号）
                    text = stripped[1:-1]
                    self._add_paragraph_with_format(text)
                else:
                    # 处理普通段落
                    text = result['text']
                    # 处理动作说明（**（...）**）格式
                    if text.strip().startswith('**（') and '）**' in text:
                        # 动作说明，保持原样但移除加粗标记
                        text = text.replace('**', '')
                        para = self.doc.add_paragraph()
                        parts = self._split_chinese_english(text)
                        for p in parts:
                            run = para.add_run(p)
                            if self._is_chinese(p):
                                run.font.name = '宋体'
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            else:
                                run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                            run.font.italic = True  # 斜体
                            run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
                        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        para.paragraph_format.line_spacing = 1.5
                    elif '**' in text:
                        # 包含加粗标记，使用加粗处理
                        self._process_bold_text(text)
                    else:
                        # 普通段落
                        self._add_paragraph_with_format(text)
        
        # 生成输出文件名
        base_name = os.path.splitext(os.path.basename(markdown_file))[0]
        date_str = datetime.now().strftime('%Y%m%d')
        output_file = f"{base_name}_{date_str}.docx"
        
        # 保存文档
        self.doc.save(output_file)
        print(f"✅ Word文档已生成：{output_file}")
        
        return output_file


def main():
    """主函数"""
    print("=" * 60)
    print("MD转Word神器")
    print("=" * 60)
    print()
    
    # 获取当前目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 查找Markdown文件
    md_files = []
    for file in os.listdir(current_dir):
        if file.endswith('.md') and not file.startswith('~'):
            md_files.append(file)
    
    if not md_files:
        print("❌ 当前目录下没有找到Markdown文件")
        print(f"当前目录：{current_dir}")
        return
    
    print(f"📁 找到 {len(md_files)} 个Markdown文件：")
    for i, file in enumerate(md_files, 1):
        print(f"  {i}. {file}")
    print()
    
    # 优先处理特定文件
    target_file = None
    priority_files = [
        'Day1_逐字稿.md',
        '三天直播课程完整逐字稿_内部渠道版_完整内容.md',
        '三天直播课程完整逐字稿_内部渠道版.md'
    ]
    
    for pf in priority_files:
        if pf in md_files:
            target_file = pf
            break
    
    if not target_file:
        # 让用户选择
        print("请选择要转换的文件（输入序号）：")
        try:
            choice = int(input().strip())
            if 1 <= choice <= len(md_files):
                target_file = md_files[choice - 1]
            else:
                print("❌ 无效的选择")
                return
        except (ValueError, KeyboardInterrupt):
            print("❌ 操作已取消")
            return
    
    if target_file:
        converter = MarkdownToWordConverter()
        output_file = converter.convert(target_file)
        
        if output_file:
            print()
            print("=" * 60)
            print("✅ 转换完成！")
            print(f"📄 输出文件：{output_file}")
            print("=" * 60)
    else:
        print("❌ 未选择文件")


if __name__ == '__main__':
    try:
        main()
    except KeyboardInterrupt:
        print("\n❌ 操作已取消")
    except Exception as e:
        print(f"\n❌ 发生错误：{e}")
        import traceback
        traceback.print_exc()

