import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# --- CONFIGURATION ---
SOURCE_DIR = r"c:\Users\温柔的男子啊\Desktop\crusor\圆心工作\RAG知识库\07_销售SOP库"
OUTPUT_PATH = r"c:\Users\温柔的男子啊\Desktop\crusor\圆心工作\销售SOP库完整文档_高级目录版.docx"

MAIN_COLOR = RGBColor(0x1E, 0x3A, 0x8A)  # Deep Blue #1E3A8A
GOLD_COLOR = RGBColor(0xD4, 0xAF, 0x37)  # Gold #D4AF37
FONT_NAME = '微软雅黑'
FONT_EN = 'Segoe UI'

class PremiumSOPConsolidator:
    def __init__(self):
        self.doc = Document()
        self._set_style()

    def _set_style(self):
        # Set margins
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(3.0)
            
        # Default font
        style = self.doc.styles['Normal']
        style.font.name = FONT_EN
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.space_after = Pt(8)

    def add_cover_page(self):
        # Add space
        for _ in range(8): self.doc.add_paragraph()
        
        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("圆心灵析 • 销售标准化体系\n")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = MAIN_COLOR
        
        # Subtitle
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("销售 SOP 库 完整汇编")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = GOLD_COLOR
        
        for _ in range(5): self.doc.add_paragraph()
        
        # Metadata
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}\n版本: V1.0 - 专业典藏版")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(100, 100, 100)
        
        self.doc.add_page_break()

    def add_toc(self, heading_list):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("详 细 目 录")
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = MAIN_COLOR
        
        self.doc.add_paragraph() # Spacer
        
        # Add Static TOC with links (simulated with text and indentation)
        for level, text in heading_list:
            if level > 2: continue # Only show H1 and H2 in static TOC to keep it clean
            p = self.doc.add_paragraph()
            if level == 1:
                run = p.add_run(text)
                run.font.bold = True
                run.font.size = Pt(12)
            else:
                p.paragraph_format.left_indent = Cm(1.0)
                run = p.add_run(f"• {text}")
                run.font.size = Pt(11)
            
        self.doc.add_page_break()
        
        # Also add the Word dynamic TOC field for professional use
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("系统索引")
        run.font.size = Pt(14)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._element.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'separate')
        run._element.append(fldChar)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        
        self.doc.add_page_break()

    def add_header_footer(self):
        for section in self.doc.sections:
            # Header
            header = section.header
            header_p = header.paragraphs[0]
            header_p.text = "圆心灵析 - 销售标准化SOP汇编"
            header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in header_p.runs:
                run.font.size = Pt(9)
                run.font.name = FONT_EN
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                run.font.color.rgb = RGBColor(150, 150, 150)
            
            # Footer with Page Numbers
            footer = section.footer
            footer_p = footer.paragraphs[0]
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Adding page number field via XML
            def add_page_number(paragraph):
                run = paragraph.add_run()
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(qn('w:fldCharType'), 'begin')
                run._element.append(fldChar)
                
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"
                run._element.append(instrText)
                
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(qn('w:fldCharType'), 'end')
                run._element.append(fldChar)

            add_page_number(footer_p)
            for run in footer_p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)

    def parse_md_to_docx(self, file_path):
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        in_code_block = False
        
        for line in lines:
            line = line.strip('\n')
            
            # Code Blocks
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                continue
            
            if in_code_block:
                p = self.doc.add_paragraph(line)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                continue

            # Headings
            if line.startswith('#'):
                level = line.count('#')
                text = line.strip('# ')
                if level == 1:
                    p = self.doc.add_heading(text, level=1)
                    for run in p.runs:
                        run.font.size = Pt(22)
                        run.font.color.rgb = MAIN_COLOR
                        run.font.name = FONT_EN
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                elif level == 2:
                    p = self.doc.add_heading(text, level=2)
                    for run in p.runs:
                        run.font.size = Pt(16)
                        run.font.color.rgb = MAIN_COLOR
                        run.font.name = FONT_EN
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                else:
                    p = self.doc.add_heading(text, level=min(level, 3))
                    for run in p.runs:
                        run.font.color.rgb = MAIN_COLOR
                continue

            # Tables (Basic support)
            if '|' in line and '-' not in line:
                # This script treats table text as normal paragraphs for safety 
                # unless we want to implement full table parsing. 
                # Given the complexity, we'll just add as special text.
                pass

            # Regular Paragraphs
            if not line.strip(): continue
            
            p = self.doc.add_paragraph()
            # Handle Bold **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.color.rgb = GOLD_COLOR
                else:
                    p.add_run(part)
            
            for run in p.runs:
                run.font.name = FONT_EN
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    def process(self):
        print("🚀 开始构建顶级 SOP 文档 (带精美目录版)...")
        
        files = [f for f in os.listdir(SOURCE_DIR) if f.endswith('.md')]
        files.sort()
        
        # Pre-scan for heading list
        heading_list = []
        for filename in files:
            path = os.path.join(SOURCE_DIR, filename)
            with open(path, 'r', encoding='utf-8') as f:
                for line in f:
                    if line.startswith('#'):
                        level = line.count('#')
                        text = line.strip('# \n')
                        heading_list.append((level, text))

        self.add_cover_page()
        self.add_toc(heading_list)
        self.add_header_footer()
        
        for i, filename in enumerate(files):
            print(f"📄 正在整合 ({i+1}/{len(files)}): {filename}")
            if i > 0:
                self.doc.add_page_break()
            
            path = os.path.join(SOURCE_DIR, filename)
            self.parse_md_to_docx(path)
            
        self.doc.save(OUTPUT_PATH)
        print(f"✨ 整合完成！输出路径: {OUTPUT_PATH}")

if __name__ == "__main__":
    consolidator = PremiumSOPConsolidator()
    consolidator.process()
