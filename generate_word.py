from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 创建文档对象
doc = Document()

# 添加标题
title = doc.add_heading('示例文档', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加副标题
doc.add_heading('这是一个使用 Python 生成的 Word 文档', level=1)

# 添加段落
doc.add_paragraph('这是文档的第一段内容。使用 python-docx 库可以很方便地生成 Word 文档。')
doc.add_paragraph('以下是一些功能特性：')

# 添加列表
features = [
    '创建和编辑文档内容',
    '设置字体样式和大小',
    '添加标题和段落',
    '创建列表和表格',
    '插入图片',
    '设置页面格式'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

# 添加表格
table = doc.add_table(rows=3, cols=3)
table.style = 'Table Grid'

# 填充表格数据
table.cell(0, 0).text = '姓名'
table.cell(0, 1).text = '年龄'
table.cell(0, 2).text = '部门'
table.cell(1, 0).text = '张三'
table.cell(1, 1).text = '28'
table.cell(1, 2).text = '技术部'
table.cell(2, 0).text = '李四'
table.cell(2, 1).text = '32'
table.cell(2, 2).text = '销售部'

# 添加段落
doc.add_paragraph('这是文档的结尾段落。您可以根据需要扩展此脚本，添加更多功能。')

# 保存文档
doc.save('example.docx')
print('Word 文档已成功生成！')
