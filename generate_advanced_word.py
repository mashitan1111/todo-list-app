from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

# 创建文档对象
doc = Document()

# 设置页面格式（A4横向）
section = doc.sections[0]
section.orientation = WD_ORIENTATION.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 创建自定义样式
styles = doc.styles

# 标题样式
title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
title_font = title_style.font
title_font.name = '微软雅黑'
title_font.size = Pt(24)
title_font.bold = True
title_font.color.rgb = RGBColor(0, 0, 128)

title_format = title_style.paragraph_format
title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format.space_after = Pt(24)

# 副标题样式
subtitle_style = styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
subtitle_font = subtitle_style.font
subtitle_font.name = '微软雅黑'
subtitle_font.size = Pt(16)
subtitle_font.italic = True
subtitle_font.color.rgb = RGBColor(128, 0, 128)

subtitle_format = subtitle_style.paragraph_format
subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_format.space_after = Pt(12)

# 正文样式
body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
body_font = body_style.font
body_font.name = '宋体'
body_font.size = Pt(12)
body_font.color.rgb = RGBColor(0, 0, 0)

body_format = body_style.paragraph_format
body_format.line_spacing = 1.5
body_format.space_after = Pt(6)
body_format.first_line_indent = Pt(24)

# 添加标题
doc.add_paragraph('高级 Word 文档生成示例', style='Custom Title')

# 添加副标题
doc.add_paragraph('使用 Python python-docx 库生成格式化文档', style='Custom Subtitle')

# 添加日期和作者
info_paragraph = doc.add_paragraph()
info_run = info_paragraph.add_run('生成日期：2026年1月5日  |  作者：AI助手')
info_run.font.size = Pt(10)
info_run.font.color.rgb = RGBColor(128, 128, 128)
info_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# 添加分隔线
doc.add_paragraph('=' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

# 第一部分：文档介绍
doc.add_heading('一、文档介绍', level=1)
doc.add_paragraph(
    '本文档展示了使用 python-docx 库生成高级 Word 文档的各种功能，' +
    '包括页面格式设置、自定义样式、表格、列表、图片插入等。',
    style='Custom Body'
)

# 第二部分：功能特性
doc.add_heading('二、主要功能特性', level=1)

# 添加带编号的列表
features = [
    '支持自定义页面格式（纸张大小、方向、边距）',
    '可以创建和应用自定义样式（字体、大小、颜色、对齐方式）',
    '支持添加各种类型的列表（项目符号、编号）',
    '能够创建和格式化表格（合并单元格、设置边框、对齐方式）',
    '支持插入图片和设置图片属性',
    '可以添加分页符和分节符',
    '支持设置页眉页脚',
    '能够添加超链接和特殊格式文本'
]

for i, feature in enumerate(features, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. ').bold = True
    p.add_run(feature)
    p.style = 'Custom Body'

# 第三部分：表格示例
doc.add_heading('三、表格示例', level=1)

doc.add_paragraph('以下是一个包含合并单元格和格式化的表格示例：', style='Custom Body')

# 创建表格 - 1行标题 + 1行表头 + 4行数据 = 6行
table = doc.add_table(rows=6, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 合并第一行单元格
header_cell = table.cell(0, 0)
header_cell.merge(table.cell(0, 3))
header_cell.text = '销售数据统计'
header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
header_cell.paragraphs[0].runs[0].bold = True
header_cell.paragraphs[0].runs[0].font.size = Pt(12)
header_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# 填充表格数据
headers = ['产品名称', '季度', '销售量(件)', '销售额(元)']
for i, header in enumerate(headers):
    cell = table.cell(1, i)
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 表格数据
table_data = [
    ['产品A', 'Q1', '1200', '60000'],
    ['产品A', 'Q2', '1500', '75000'],
    ['产品B', 'Q1', '800', '40000'],
    ['产品B', 'Q2', '1000', '50000']
]

# 从第2行开始填充数据 (0-based index)
for row_idx, row_data in enumerate(table_data):
    for col_idx, cell_data in enumerate(row_data):
        cell = table.cell(row_idx + 2, col_idx)
        cell.text = cell_data
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if col_idx > 1:  # 数字列右对齐
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 添加分页符
doc.add_page_break()

# 第四部分：图片示例
doc.add_heading('四、图片示例', level=1)
doc.add_paragraph('以下是插入图片的示例（如果有图片文件可以替换路径）：', style='Custom Body')

# 添加文本框（使用占位符）
p = doc.add_paragraph()
p.add_run('注意：').bold = True
p.add_run(' 由于没有实际图片文件，此示例使用文字说明。在实际使用中，可以使用以下代码插入图片：')
p.add_run('\n\ndoc.add_picture("image.jpg", width=Inches(4.0))').font.name = 'Courier New'
p.style = 'Custom Body'

# 第五部分：特殊格式文本
doc.add_heading('五、特殊格式文本', level=1)

# 加粗文本
p = doc.add_paragraph()
p.add_run('这是加粗文本。').bold = True
p.style = 'Custom Body'

# 斜体文本
p = doc.add_paragraph()
p.add_run('这是斜体文本。').italic = True
p.style = 'Custom Body'

# 下划线文本
p = doc.add_paragraph()
p.add_run('这是下划线文本。').underline = True
p.style = 'Custom Body'

# 带颜色的文本
p = doc.add_paragraph()
p.add_run('这是红色文本。').font.color.rgb = RGBColor(255, 0, 0)
p.style = 'Custom Body'

# 高亮文本
p = doc.add_paragraph()
p.add_run('这是高亮文本。').font.highlight_color = WD_COLOR_INDEX.YELLOW
p.style = 'Custom Body'

# 超链接
doc.add_heading('六、超链接', level=1)
p = doc.add_paragraph()
p.add_run('python-docx 官方文档：').bold = True
# 使用文本形式表示超链接（python-docx 中的超链接需要使用 XML 操作添加）
p.add_run(' https://python-docx.readthedocs.io/')
p.style = 'Custom Body'

# 第七部分：总结
doc.add_heading('七、总结', level=1)
doc.add_paragraph(
    '通过本示例，您可以了解到使用 python-docx 库生成 Word 文档的强大功能。' +
    '您可以根据实际需求扩展这些功能，创建更加专业和精美的文档。',
    style='Custom Body'
)

# 添加页脚
for section in doc.sections:
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.text = '第 {PAGE} 页，共 {NUMPAGES} 页'
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置页脚字体大小
    if footer_paragraph.runs:
        footer_paragraph.runs[0].font.size = Pt(9)

# 保存文档
doc.save('advanced_example.docx')
print('高级 Word 文档已成功生成！')
